from docx import Document
import random


def generate_recommendations(num_recommendations):
    recommendations = [
        "Прежде чем сесть за руль, особенно если Вы в нетрезвом состоянии, оцените свои способности и учитывайте погодные условия, состояние дороги и уровень освещения! Если нужно, дайте кому-то другому управлять автомобилем.",
        "Технические неисправности, особенно в сочетании с нетрезвым вождением, могут стать причиной серьезных аварий, поэтому регулярно проверяйте состояние вашего автомобиля и не откладывайте ремонт, если что-то выходит из строя!",
        "Никогда не садитесь за руль, находясь в нетрезвом состоянии, и всегда используйте ремень безопасности. Не рискуйте своей жизнью и жизнями других людей!",
        "Если Вы планируете употребить алкоголь или наркотические препараты, убедитесь, что у Вас есть «план B» для возвращения домой. Используйте такси, общественный транспорт или обратитесь к друзьям или семье за помощью!",
        "Езда на мотоциклах и подобных транспортных средствах является опасной даже в трезвом состоянии, а в нетрезвом становится еще более рискованной, поэтому исключите поездку за рулем под воздействием алкоголя или наркотиков!",
        "Если Вы владелец автомобиля NISSAN Sunny или HONDA CR-V, то будьте бдительнее за рулем! Постарайтесь соблюдать правила дорожного движения и поддерживать контроль над собой, особенно если у Вас есть склонность к нетрезвому вождению!",
        "Регулярно проверяйте техническое состояние вашего автомобиля и убедитесь, что все системы работают исправно, поскольку технические неисправности могут стать причиной ДТП, особенно если Вы в нетрезвом состоянии!",
        "Даже небольшое количество алкоголя или наркотиков может существенно повлиять на вашу реакцию и восприятие, если Вы употребляли эти вещества, то лучше воздержитесь от управления автомобилем!",
        "Если все же Вы решаетесь сесть за руль, будьте особенно осторожны и бдительны на дороге. Помните, что ваша реакция и способность принимать решения могут быть снижены, поэтому сосредоточьтесь на безопасном вождении!"
    ]

    # Выбираем случайные рекомендации
    selected_recommendations = random.sample(recommendations, num_recommendations)

    # Добавляем рекомендацию о серьезных последствиях в конец списка
    selected_recommendations.append(
        "Помните, что Ваше решение сесть за руль в нетрезвом состоянии может иметь серьезные последствия для Вас и для других! Подумайте о том, какие риски вы готовы принять, прежде чем принимать это решение!")

    return selected_recommendations


# Генерируем 5 случайных рекомендаций
num_recommendations = 5
generated_recommendations = generate_recommendations(num_recommendations)

for recommendation in generated_recommendations:
    print(recommendation)

# Создаем новый документ MS Word
doc = Document()

# Добавляем заголовок
doc.add_heading('Рекомендации', level=1)

# Добавляем рекомендации в документ
for idx, recommendation in enumerate(generated_recommendations, start=1):
    doc.add_paragraph(f'{idx}. {recommendation}')

# Сохраняем документ в файл
doc.save('рекомендации.docx')

